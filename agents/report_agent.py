from __future__ import annotations
from dataclasses import dataclass
import os
import logging
import asyncio
import json

from pydantic import BaseModel
from datetime import datetime
from pydantic_ai import Agent, RunContext
from pydantic_ai.models.openai import OpenAIModel
from openai import AsyncOpenAI
from docx import Document
from openai import AsyncOpenAI
from config.config import settings
from utils.ppt_report import generate_ppt_report, prepare_complete_placeholders

class ReportDeps(BaseModel):
    openai_client: AsyncOpenAI

    class Config:
        arbitrary_types_allowed = True

# Configurar el modelo LLM a utilizar
llm = settings.llm_model
model = OpenAIModel(llm, api_key=settings.openai_api_key)

openai_client_instance = AsyncOpenAI(api_key=settings.openai_api_key)

# Configuración básica de logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

document = Document()

def save_report_as_word(report_text: str, file_path: str):
    """
    Convierte un texto formateado en un documento de Word (.docx) y lo guarda en la ruta indicada.
    """
      
    # Separa el contenido para identificar encabezados y párrafos
    for line in report_text.splitlines():
        stripped_line = line.strip()
        if not stripped_line:
            continue  # Ignora líneas vacías
        if stripped_line.startswith("## "):
            document.add_heading(stripped_line.replace("## ", ""), level=2)
        elif stripped_line.startswith("# "):
            document.add_heading(stripped_line.replace("# ", ""), level=1)
        else:
            document.add_paragraph(stripped_line)
    
    document.save(file_path)


# Definir el prompt de sistema para el Agente de Report
report_system_prompt = """
Eres un generador de informes especializado. Tu tarea es crear un report claro, conciso y bien estructurado basado en el output de un Agente de Compliance. 

El informe debe incluir las siguientes secciones:
1. Executive Summary
2. Findings
3. Recommendations
4. Conclusion

Si se proporciona una plantilla (template), utiliza los placeholders correspondientes para cada sección.
Utiliza un formato Markdown para el documento final.

**Nota:** Esta herramienta se utiliza para generar informes en formato Word. Si se requiere un informe en formato PPT, el agente de compliance deberá haber respondido en formato JSON, y se utilizará una herramienta específica para PPT.

"""

report_agent = Agent(
    model=model,
    system_prompt=report_system_prompt,
    deps_type=ReportDeps,
    retries=2
)

@report_agent.tool
async def create_compliance_report(ctx: RunContext[ReportDeps], compliance_output: str, template: str = None) -> str:
    """
    Genera un informe de compliance utilizando el output del Agente de Compliance y lo guarda como un archivo de Word.
    Retorna la ruta del archivo generado.
    """
    # Si no se proporciona una plantilla, se utiliza una por defecto (en Markdown)
    if template is None:
        template = (
            "# Compliance Report\n\n"
            "## Executive Summary\n"
            "{executive_summary}\n\n"
            "## Findings\n"
            "{findings}\n\n"
            "## Recommendations\n"
            "{recommendations}\n\n"
            "## Conclusion\n"
            "{conclusion}\n"
        )
    
    # Crear el prompt combinando el output y la plantilla
    prompt = (
        f"Con base en el siguiente análisis de compliance:\n\n{compliance_output}\n\n"
        "Genera UN INFORME FINAL y DETALLADO que incluya las secciones: Executive Summary, Findings, Recommendations y Conclusion. "
    "DEBERÁS devolver únicamente el informe final, sin solicitar acciones adicionales, sin invocar herramientas, ni generar instrucciones para continuar. "
    "Utiliza la siguiente plantilla y reemplaza cada placeholder con la información correspondiente:\n\n"
    f"{template}\n"
    )

    # Llamada al modelo (internamente, el agente genera el informe en formato Markdown)
    response = await report_agent.run(
        prompt, 
        deps=ReportDeps(openai_client=ctx.deps.openai_client),
        usage=ctx.usage)
    report_text = response.data
    
    # Asegurarse de que la carpeta output exista
    output_dir = os.path.join(os.getcwd(), "output")
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    # Definir el nombre del archivo con marca de tiempo
    file_name = f"compliance_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    output_file = os.path.join(output_dir, file_name)
    
    # Guardar el informe como Word
    save_report_as_word(report_text, output_file)
    
    return f"Informe generado y guardado en: {output_file}"



