# agents/report_agent.py - Versión mejorada con soporte para templates
from __future__ import annotations as _annotations

import os
import datetime
import logging
import re
from typing import Optional, Dict, Any, List

from pydantic import BaseModel
from pydantic_ai import Agent, RunContext
from pydantic_ai.models.openai import OpenAIModel
from openai import AsyncOpenAI
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.core.config import settings
from utils.ppt_report import generate_ppt_report, prepare_complete_placeholders

# Configuración de logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Inicialización del modelo
llm = settings.llm_model
model = OpenAIModel(llm, api_key=settings.openai_api_key)

class ReportDeps(BaseModel):
    output_folder: str = "output/reports"
    template_folder: str = "agents/templates"  # Nueva carpeta para templates
    openai_client: AsyncOpenAI

    class Config:
        arbitrary_types_allowed = True

class ReportResult(BaseModel):
    file_path: str
    report_type: str = "word"
    message: str = "Reporte generado exitosamente"
    sections_filled: List[str] = []  # Nuevas secciones completadas

# Mapeo de placeholders del template
TEMPLATE_PLACEHOLDERS = {
    "{{EXECUTIVE_SUMMARY}}": "executive_summary",
    "{{ALCANCE}}": "scope", 
    "{{FINDINGS}}": "findings",
    "{{CONCLUSIONES_RECOMENDACIONES}}": "conclusions_recommendations"
}

report_system_prompt = """
Eres un agente experto en generación de reportes normativos sobre regulaciones financieras y no financieras. 
Tu tarea es completar templates de informes estructurados con análisis detallados y recomendaciones específicas.

Para cada sección del template debes:
1. Analizar la información regulatoria proporcionada
2. Generar contenido específico, detallado y profesional
3. Mantener coherencia entre todas las secciones
4. Usar terminología técnica apropiada del sector financiero
5. Proporcionar recomendaciones accionables

El informe debe ser comprehensivo y estar listo para revisión por parte de equipos de compliance.
"""

report_agent = Agent(
    model=model,
    system_prompt=report_system_prompt,
    deps_type=ReportDeps,
    result_type=ReportResult,
    retries=2
)

def find_placeholders_in_document(doc_path: str) -> List[str]:
    """
    Encuentra todos los placeholders en un documento Word.
    
    Args:
        doc_path: Ruta al documento template
        
    Returns:
        Lista de placeholders encontrados
    """
    doc = Document(doc_path)
    placeholders = []
    
    # Buscar en párrafos
    for paragraph in doc.paragraphs:
        text = paragraph.text
        # Buscar patrones que empiecen con {{ y buscar el cierre }}
        # Esto ayudará a encontrar incluso placeholders mal formateados
        found = re.findall(r'\{\{[^}]+\}\}', text)
        placeholders.extend(found)
        
        # También buscar placeholders con errores tipográficos conocidos
        if "{{Executive Summary" in text:
            placeholders.append("{{Executive Summary}}")
    
    # Buscar en tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = paragraph.text
                    found = re.findall(r'\{\{[^}]+\}\}', text)
                    placeholders.extend(found)
    
    return list(set(placeholders))

def replace_placeholder_in_document(doc: Document, placeholder: str, content: str):
    """
    Reemplaza un placeholder específico en todo el documento con el contenido generado.
    
    Args:
        doc: Documento de Word
        placeholder: Placeholder a reemplazar (ej: "{{EXECUTIVE_SUMMARY}}")
        content: Contenido que reemplazará el placeholder
    """
    # Reemplazar en párrafos
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            # Limpiar el párrafo y agregar el nuevo contenido
            paragraph.clear()
            paragraph.add_run(content)
    
    # Reemplazar en tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if placeholder in paragraph.text:
                        paragraph.clear()
                        paragraph.add_run(content)

@report_agent.tool
async def generate_report_from_template(
    ctx: RunContext[ReportDeps], 
    analysis_data: str, 
    template_name: str = "Template_Regulatory_Report_AgentIA.docx",
    regulation_name: str = "Normativa Analizada",
    output_filename: str = None
) -> ReportResult:
    """
    Genera un reporte usando un template de Word existente.
    
    Args:
        analysis_data: Datos del análisis regulatorio
        template_name: Nombre del archivo template
        regulation_name: Nombre de la regulación analizada
        output_filename: Nombre del archivo de salida (opcional)
        
    Returns:
        ReportResult con información del archivo generado
    """
    try:
        # Construir rutas
        template_path = os.path.join(ctx.deps.template_folder, template_name)
        
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template no encontrado: {template_path}")
        
        # Cargar el template
        doc = Document(template_path)
        logger.info(f"Template cargado exitosamente desde: {template_path}")
        
        # Encontrar placeholders en el documento
        placeholders = find_placeholders_in_document(template_path)
        logger.info(f"Placeholders encontrados: {placeholders}")
        
        # Reemplazar fecha y nombre de regulación
        current_date = datetime.datetime.now().strftime("%d/%m/%Y")
        replace_placeholder_in_document(doc, "[Fecha]", current_date)
        replace_placeholder_in_document(doc, "[Ley analizada]", regulation_name)
        
        # Generar contenido para cada sección
        sections_filled = []
        
        for placeholder in placeholders:
            if placeholder in TEMPLATE_PLACEHOLDERS:
                section_key = TEMPLATE_PLACEHOLDERS[placeholder]
                content = await generate_section_content(
                    ctx.deps.openai_client, 
                    section_key, 
                    analysis_data, 
                    regulation_name
                )
                
                replace_placeholder_in_document(doc, placeholder, content)
                sections_filled.append(section_key)
                logger.info(f"Sección completada: {section_key}")
        
        # Guardar el documento
        if not os.path.exists(ctx.deps.output_folder):
            os.makedirs(ctx.deps.output_folder)
        
        if output_filename is None:
            timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
            output_filename = f"Analisis_Regulatorio_{regulation_name.replace(' ', '_')}_{timestamp}.docx"
        
        output_path = os.path.join(ctx.deps.output_folder, output_filename)
        doc.save(output_path)
        
        logger.info(f"Reporte generado exitosamente: {output_path}")
        
        return ReportResult(
            file_path=output_path,
            report_type="word",
            message=f"Reporte generado usando template {template_name}",
            sections_filled=sections_filled
        )
        
    except Exception as e:
        logger.error(f"Error generando reporte desde template: {e}")
        raise e

async def generate_section_content(
    openai_client: AsyncOpenAI, 
    section_key: str, 
    analysis_data: str, 
    regulation_name: str
) -> str:
    """
    Genera contenido específico para una sección del reporte.
    
    Args:
        openai_client: Cliente de OpenAI
        section_key: Clave de la sección a generar
        analysis_data: Datos del análisis regulatorio
        regulation_name: Nombre de la regulación
        
    Returns:
        Contenido generado para la sección
    """
    
    section_prompts = {
        "executive_summary": f"""
Basado en el análisis de la {regulation_name}:

{analysis_data}

Genera un EXECUTIVE SUMMARY profesional de 4-6 párrafos que incluya:
1. Propósito y alcance de la regulación
2. Principales hallazgos del análisis
3. Impacto en las operaciones bancarias/financieras
4. Conclusiones clave sobre el cumplimiento requerido

Debe ser conciso pero comprehensivo, dirigido a alta dirección.
""",
        
        "scope": f"""
Basado en el análisis de la {regulation_name}:

{analysis_data}

Define el ALCANCE de este informe regulatorio especificando:
1. Objetivos específicos del análisis
2. Normativas y regulaciones cubiertas
3. Áreas organizacionales afectadas
4. Limitaciones del análisis
5. Metodología utilizada

Mantén un tono profesional y específico.
""",
                
        "findings": f"""
Basado en el análisis de la {regulation_name}:

{analysis_data}

Desarrolla una sección de FINDINGS (HALLAZGOS) que destaque:
1. Los hallazgos más relevantes y críticos de la regulación
2. Aspectos clave que impactan directamente a las entidades reguladas
3. Nuevos requisitos o cambios significativos introducidos
4. Elementos de la normativa que requieren atención prioritaria
5. Interpretaciones importantes de artículos o disposiciones específicas
6. Implicaciones prácticas para el cumplimiento normativo
7. Puntos de especial vigilancia regulatoria

Estructura los findings de forma clara, priorizando por nivel de impacto e importancia.
""",
        
        "conclusions_recommendations": f"""
Basado en el análisis de la {regulation_name}:

{analysis_data}

Desarrolla CONCLUSIONES Y RECOMENDACIONES que incluyan:
1. Conclusiones principales del análisis regulatorio
2. Recomendaciones estratégicas para la alta dirección
3. Recomendaciones operativas específicas
4. Consideraciones de implementación
5. Próximos pasos sugeridos
6. Recomendaciones para monitoreo continuo

Las recomendaciones deben ser específicas, accionables y priorizadas.
"""
    }
    
    prompt = section_prompts.get(section_key, f"Genera contenido para la sección {section_key} basado en: {analysis_data}")
    
    try:
        completion = await openai_client.chat.completions.create(
            model=settings.llm_model,
            temperature=0.1,
            max_tokens=5000,
            messages=[
                {
                    "role": "system", 
                    "content": "Eres un experto en análisis regulatorio financiero. Genera contenido profesional, detallado y específico para informes de compliance."
                },
                {"role": "user", "content": prompt}
            ]
        )
        
        return completion.choices[0].message.content.strip()
        
    except Exception as e:
        logger.error(f"Error generando contenido para sección {section_key}: {e}")
        return f"[Error generando contenido para {section_key}]"

# Función principal para procesar consultas con template
async def process_report_query(
    query: str, 
    analysis_data: str, 
    deps: ReportDeps,
    template_name: str = "Template_Regulatory_Report_AgentIA.docx",
    regulation_name: str = None
) -> ReportResult:
    """
    Procesa una consulta para generar un reporte usando un template específico.
    
    Args:
        query: Consulta del usuario
        analysis_data: Datos del análisis regulatorio
        deps: Dependencias del agente
        template_name: Nombre del template a usar
        regulation_name: Nombre de la regulación (se extrae de query si no se proporciona)
        
    Returns:
        ReportResult con información del reporte generado
    """
    logger.info(f"Procesando consulta con template: {query[:100]}...")
    
    # Extraer nombre de regulación si no se proporciona
    if regulation_name is None:
        try:
            completion = await deps.openai_client.chat.completions.create(
                model=settings.llm_model,
                temperature=0.1,
                messages=[
                    {
                        "role": "system", 
                        "content": "Extrae el nombre de la regulación o ley mencionada en la consulta. Responde solo con el nombre."
                    },
                    {"role": "user", "content": query}
                ]
            )
            regulation_name = completion.choices[0].message.content.strip()
        except:
            regulation_name = "Regulación Analizada"
    
    # Usar el nuevo tool para generar el reporte
    try:
        result = await report_agent.run(
            f"Genera un reporte regulatorio completo para: {query}",
            deps=deps,
            message_history=[],
        )
        
        # Llamar directamente a la función de template
        return await generate_report_from_template(
            RunContext(deps=deps, retry=0),
            analysis_data=analysis_data,
            template_name=template_name,
            regulation_name=regulation_name
        )
        
    except Exception as e:
        logger.error(f"Error procesando consulta con template: {e}")
        return ReportResult(
            file_path="",
            message=f"Error generando reporte con template: {str(e)}"
        )